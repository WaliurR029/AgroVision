from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Title
title = doc.add_heading('🌿 AgroVision - Complete Project Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============ PART A ============
doc.add_heading('PART A: PROJECT OVERVIEW & CONCEPTS', level=1)

doc.add_heading('📁 Project Structure Overview', level=2)
doc.add_paragraph('''
AgroVision/
├── AgroVision.ipynb          # Training notebook (model creation)
├── app.py                    # Streamlit web application
├── predict.py                # Command-line prediction script
├── best_agrovision_model.keras  # Trained model file
├── class_indices.json        # Class label mappings
├── PlantVillage/             # Dataset folder (15 classes)
├── requirements.txt          # Python dependencies
├── .streamlit/config.toml    # Streamlit configuration
├── accuracy_curve.png        # Training accuracy graph
├── loss_curve.png            # Training loss graph
└── confusion_matrix.png      # Model evaluation matrix
''')

doc.add_heading('🎯 1. PROJECT GOAL', level=2)
doc.add_paragraph('Problem: Farmers often can\'t identify plant diseases early, leading to crop loss.')
doc.add_paragraph('''Solution: An AI system that:
1. Takes a photo of a plant leaf
2. Identifies if it's healthy or diseased
3. Tells which disease it has
4. Provides treatment recommendations''')

doc.add_paragraph('Supported Crops: Tomato, Potato, Pepper (Bell)')

doc.add_heading('15 Classes:', level=3)
classes = '''1. Pepper__bell___Bacterial_spot - Disease
2. Pepper__bell___healthy - Healthy
3. Potato___Early_blight - Disease
4. Potato___Late_blight - Disease
5. Potato___healthy - Healthy
6. Tomato_Bacterial_spot - Disease
7. Tomato_Early_blight - Disease
8. Tomato_Late_blight - Disease
9. Tomato_Leaf_Mold - Disease
10. Tomato_Septoria_leaf_spot - Disease
11. Tomato_Spider_mites - Pest
12. Tomato__Target_Spot - Disease
13. Tomato__Tomato_YellowLeaf__Curl_Virus - Virus
14. Tomato__Tomato_mosaic_virus - Virus
15. Tomato_healthy - Healthy'''
doc.add_paragraph(classes)

doc.add_heading('🧠 2. THE NEURAL NETWORK ARCHITECTURE', level=2)
doc.add_heading('What is MobileNetV2?', level=3)
doc.add_paragraph('MobileNetV2 is a pre-trained Convolutional Neural Network (CNN) developed by Google, trained on ImageNet (1.4 million images, 1000 classes).')

doc.add_paragraph('''Why MobileNetV2?
• Lightweight: Only ~3.4 million parameters (vs. 138M in VGG16)
• Fast: Designed for mobile devices
• Accurate: Uses "inverted residuals" and "linear bottlenecks"
• Transfer Learning: Already knows how to recognize edges, textures, shapes''')

doc.add_heading('Model Architecture:', level=3)
doc.add_paragraph('''
INPUT IMAGE (256×256×3) - RGB image, 256 pixels × 256 pixels
        ↓
MobileNetV2 (Base) - Pre-trained feature extractor, Extracts 1280 features
        ↓
GlobalAveragePooling2D() - Reduces 8×8×1280 → 1280 values
        ↓
Dropout(0.4) - Randomly drops 40% neurons (prevents overfitting)
        ↓
Dense(256, activation='relu') - Fully connected layer, 256 neurons
        ↓
Dropout(0.3) - Drops 30% neurons
        ↓
Dense(15, activation='softmax') - Output layer: 15 class probabilities
''')

doc.add_heading('Key Concepts:', level=3)
concepts = '''
• Transfer Learning: Using a model trained on one task (ImageNet) for another task (plant diseases). The model already knows basic features like edges, colors, textures.

• Fine-tuning: base_model.trainable = True means we allow MobileNetV2's weights to be updated during training, not just the top layers.

• GlobalAveragePooling2D: Instead of flattening (which creates too many parameters), it takes the average of each feature map. Reduces overfitting.

• Dropout: During training, randomly "turns off" neurons. Forces the network to not rely on any single neuron. Prevents overfitting.

• ReLU: Rectified Linear Unit: f(x) = max(0, x). Introduces non-linearity. Fast to compute.

• Softmax: Converts raw scores to probabilities that sum to 1.
'''
doc.add_paragraph(concepts)

doc.add_heading('📊 3. DATA PREPROCESSING & AUGMENTATION', level=2)
doc.add_paragraph('''ImageDataGenerator Parameters:

datagen = ImageDataGenerator(
    rescale=1./255,              # Normalize pixels from 0-255 → 0-1
    validation_split=0.2,        # 80% train, 20% validation
    rotation_range=20,           # Rotate images ±20 degrees
    zoom_range=0.15,             # Zoom in/out by 15%
    brightness_range=[0.8, 1.2], # Vary brightness 80%-120%
    horizontal_flip=True,        # Flip images horizontally
    fill_mode='nearest'          # Fill empty pixels with nearest value
)''')

doc.add_heading('Why Data Augmentation?', level=3)
doc.add_paragraph('''Problem: Limited training data leads to overfitting.
Solution: Artificially create more training samples by transforming existing images.

• Rotation - Leaves can be photographed at any angle
• Zoom - Camera distance varies
• Brightness - Lighting conditions vary (sunny/cloudy)
• Horizontal Flip - Leaves can face either direction''')

doc.add_heading('⚖️ 4. CLASS WEIGHTS (Handling Imbalanced Data)', level=2)
doc.add_paragraph('''Why Class Weights?

Problem: Some classes have more images than others. The model might learn to always predict the majority class.

Solution: Give higher weight to minority classes during training.

Formula: w_j = n_samples / (n_classes × n_samples_j)

Example:
• If "Potato___healthy" has only 100 images
• But "Tomato_healthy" has 500 images
• Potato___healthy gets weight = 5× higher''')

doc.add_heading('🔧 5. OPTIMIZER & LOSS FUNCTION', level=2)
doc.add_heading('AdamW Optimizer', level=3)
doc.add_paragraph('''optimizer = tf.keras.optimizers.AdamW(
    learning_rate=5e-5,    # 0.00005 - very small for fine-tuning
    weight_decay=1e-5      # L2 regularization
)

Adam = Adaptive Moment Estimation
• Combines momentum + RMSprop
• Adapts learning rate for each parameter
• AdamW adds decoupled weight decay (better regularization)

Why small learning rate (5e-5)?
• MobileNetV2 is already trained
• We want to make small adjustments, not destroy learned features
• Large LR would cause "catastrophic forgetting"''')

doc.add_heading('Loss Function: Categorical Cross-Entropy with Label Smoothing', level=3)
doc.add_paragraph('''loss=tf.keras.losses.CategoricalCrossentropy(label_smoothing=0.1)

Label Smoothing (0.1):
• Instead of hard labels [0, 0, 1, 0, ...]
• Use soft labels [0.0067, 0.0067, 0.9, 0.0067, ...]
• Prevents overconfidence, improves generalization''')

doc.add_heading('📈 6. CALLBACKS (Training Control)', level=2)
doc.add_paragraph('''EarlyStopping:
EarlyStopping(monitor='val_accuracy', patience=15, restore_best_weights=True)
• Monitors: Validation accuracy
• Patience: Waits 15 epochs without improvement before stopping
• Restore: Returns to the best weights found

ReduceLROnPlateau:
ReduceLROnPlateau(monitor='val_loss', factor=0.2, patience=5)
• If validation loss doesn't improve for 5 epochs
• Multiply learning rate by 0.2 (reduce it)
• Helps escape local minima

ModelCheckpoint:
ModelCheckpoint("best_agrovision_model.keras", monitor="val_accuracy", save_best_only=True)
• Saves model only when validation accuracy improves
• You always have the best model saved''')

doc.add_heading('🏋️ 7. TRAINING PROCESS', level=2)
doc.add_paragraph('''history = model.fit(
    train_data,
    validation_data=val_data,
    epochs=15,
    callbacks=[early_stop, reduce_lr, checkpoint],
    class_weight=class_weights
)

What Happens Each Epoch:
1. Batch 1: Process 32 images → Calculate loss → Backpropagate → Update weights
2. Batch 2: Process 32 images → Calculate loss → Backpropagate → Update weights
3. ... (all batches)
4. Training complete: Calculate average training accuracy & loss
5. Validation: Test on validation set (no weight updates)
6. Callbacks check: Should we stop? Reduce LR? Save model?
7. Move to next Epoch''')

doc.add_heading('❓ COMMON PROFESSOR QUESTIONS', level=2)

qa = [
    ("Q1: Why MobileNetV2 and not VGG16 or ResNet?", "MobileNetV2 is lightweight (3.4M vs 138M parameters), fast inference (good for real-time apps), good accuracy despite small size, designed for edge devices (phones, embedded systems)."),
    ("Q2: What is Transfer Learning?", "Using knowledge from one task (ImageNet - 1000 classes) to solve another task (plant diseases - 15 classes). The pre-trained model already understands basic image features (edges, textures, shapes). We just teach it plant-specific features."),
    ("Q3: Why 256×256 image size?", "Balance between detail preservation (larger = more detail), memory usage (larger = more GPU memory), training speed (smaller = faster). MobileNetV2 was trained on 224×224, so similar sizes work well."),
    ("Q4: What if the model predicts wrong?", "We show confidence score. If < 60%, we display 'Inconclusive'. The model also shows top 3 predictions so users can see alternatives."),
    ("Q5: How does Dropout prevent overfitting?", "During training, randomly 'turns off' neurons (sets output to 0). This prevents neurons from co-adapting, forces network to learn redundant representations, acts like training multiple smaller networks (ensemble effect)."),
    ("Q6: What's the difference between Training and Validation accuracy?", "Training accuracy: Performance on data the model learns from. Validation accuracy: Performance on unseen data (true test). If training >> validation → Overfitting. If both are similar → Good generalization."),
    ("Q7: Why use Softmax in the output layer?", "Softmax converts raw scores to probabilities: All outputs sum to 1.0 (100%), each output is between 0 and 1, highest probability = predicted class."),
    ("Q8: What is Backpropagation?", "Algorithm to update weights: 1) Forward pass: Input → Prediction, 2) Calculate loss (error), 3) Backward pass: Calculate gradients, 4) Update weights: weight = weight - learning_rate × gradient."),
    ("Q9: Why save as .keras instead of .h5?", ".keras is TensorFlow 2.x native format: Saves model architecture + weights + optimizer state, better compatibility with new TensorFlow versions, recommended by Keras team."),
    ("Q10: What real-world impact does this have?", "Early disease detection saves crops, farmers don't need expert knowledge, reduces pesticide misuse (targeted treatment), can work offline on mobile devices, scales to millions of farmers."),
]

for q, a in qa:
    doc.add_paragraph(q, style='Intense Quote')
    doc.add_paragraph(a)

# ============ PART B ============
doc.add_page_break()
doc.add_heading('PART B: COMPLETE CODE EXPLANATION (Line by Line)', level=1)

doc.add_heading('📓 SECTION 1: AgroVision.ipynb (Training Notebook)', level=2)

doc.add_heading('Cell 1: Core Imports', level=3)
doc.add_paragraph('''# Core imports
import os, json, pickle
import numpy as np
import tensorflow as tf
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications import MobileNetV2
from tensorflow.keras.models import Sequential, load_model
from tensorflow.keras.layers import Dense, Dropout, GlobalAveragePooling2D
from tensorflow.keras.callbacks import EarlyStopping, ReduceLROnPlateau, ModelCheckpoint
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.utils.class_weight import compute_class_weight
from sklearn.metrics import confusion_matrix, classification_report''')

doc.add_paragraph('''Line-by-line explanation:
• os - Operating system operations (file paths, directories)
• json - Read/write JSON files (for class_indices.json)
• pickle - Serialize Python objects (save training history)
• numpy as np - Numerical computing - arrays, math operations
• tensorflow as tf - Google's deep learning framework
• ImageDataGenerator - Loads images from folders + applies augmentation
• MobileNetV2 - Pre-trained CNN model (our base model)
• Sequential - Stack layers linearly (layer1 → layer2 → layer3)
• load_model - Load saved .keras model files
• Dense - Fully connected neural network layer
• Dropout - Regularization layer (randomly drops neurons)
• GlobalAveragePooling2D - Reduces spatial dimensions by averaging
• EarlyStopping - Stop training when no improvement
• ReduceLROnPlateau - Reduce learning rate when stuck
• ModelCheckpoint - Save best model during training
• matplotlib.pyplot - Plotting graphs
• seaborn - Statistical visualization (prettier plots)
• compute_class_weight - Calculate weights for imbalanced classes
• confusion_matrix - Shows prediction errors in matrix form
• classification_report - Precision, recall, F1-score per class''')

doc.add_heading('Cell 2: Dataset Path', level=3)
doc.add_paragraph('''dataset_path = r"C:\\Users\\...\\PlantVillage"

• r"..." - Raw string - backslashes are treated literally (no escape sequences)
• PlantVillage/ - Contains 15 subfolders, each subfolder = one class

Folder Structure:
PlantVillage/
├── Pepper__bell___Bacterial_spot/    # ~1000 images
├── Pepper__bell___healthy/           # ~1000 images
├── ... (15 folders total)

Why this structure? flow_from_directory() automatically:
1. Reads folder names as class labels
2. Assigns numeric indices alphabetically
3. Loads all images from each folder''')

doc.add_heading('Cell 3: Data Generators', level=3)
doc.add_paragraph('''datagen = ImageDataGenerator(
    rescale=1./255,           # Normalization: 0-255 → 0-1
    validation_split=0.2,     # 80% train, 20% validation
    rotation_range=20,        # ±20° rotation
    zoom_range=0.15,          # ±15% zoom
    brightness_range=[0.8, 1.2],  # 80%-120% brightness
    horizontal_flip=True,     # Random horizontal flip
    fill_mode='nearest'       # Fill empty pixels
)

train_data = datagen.flow_from_directory(
    dataset_path,
    target_size=(256, 256),   # Resize all images to 256×256
    batch_size=32,            # Process 32 images at once
    class_mode='categorical', # One-hot encoding
    subset='training',        # Use 80% for training
    seed=42                   # Reproducibility
)''')

doc.add_heading('Cell 5: Model Building & Training', level=3)
doc.add_paragraph('''Part A: Class Weights
class_weights = compute_class_weight('balanced', ...)
• Calculates weights inversely proportional to class frequency
• cap at 5.0 to prevent instability

Part B: Base Model
base_model = MobileNetV2(weights='imagenet', include_top=False, input_shape=(256, 256, 3))
base_model.trainable = True
• weights='imagenet' - Load pre-trained weights
• include_top=False - Remove original 1000-class output layer
• trainable=True - Allow fine-tuning all layers

Part C: Custom Classifier
model = Sequential([
    base_model,                    # Feature extraction
    GlobalAveragePooling2D(),      # Reduce dimensions
    Dropout(0.4),                  # Regularization
    Dense(256, activation='relu'), # Hidden layer
    Dropout(0.3),                  # More regularization
    Dense(15, activation='softmax') # Output: 15 classes
])

Part D: Compilation
optimizer = tf.keras.optimizers.AdamW(learning_rate=5e-5, weight_decay=1e-5)
model.compile(optimizer=optimizer, 
              loss=CategoricalCrossentropy(label_smoothing=0.1),
              metrics=['accuracy'])

Part E: Training
history = model.fit(train_data, validation_data=val_data, epochs=15, 
                    callbacks=[early_stop, reduce_lr, checkpoint],
                    class_weight=class_weights)''')

doc.add_heading('🌐 SECTION 2: app.py (Streamlit Web App)', level=2)

doc.add_paragraph('''Key Components:

1. Caching for Performance:
@st.cache_resource
def load_ai_model():
    return load_model("best_agrovision_model.keras")
• Load model only once, reuse for all predictions

2. Session State for History:
if 'history' not in st.session_state:
    st.session_state['history'] = []
• Persists data between Streamlit reruns

3. Prediction Logic:
img = image_pil.resize((256, 256))        # Resize to model input
img_array = image.img_to_array(img)        # PIL → NumPy
img_array = np.expand_dims(img_array, 0)   # Add batch dimension
img_array /= 255.0                         # Normalize

predictions = model.predict(img_array)[0]  # Get probabilities
predicted_index = np.argmax(predictions)   # Highest probability
confidence = predictions[predicted_index]  # Confidence score

4. Display Results:
if confidence > 0.6:
    st.success(f"Detected: {disease_name}")
    st.info(disease_solutions[predicted_class])
else:
    st.warning("Analysis Inconclusive")''')

doc.add_heading('🔑 CRITICAL CONCEPTS SUMMARY', level=2)
doc.add_paragraph('''
• Transfer Learning - Using pre-trained MobileNetV2, don't need millions of images
• Fine-tuning - Unfreezing base model layers, adapts ImageNet features to plants
• Data Augmentation - Artificial image variations, prevents overfitting
• Class Weights - Balance imbalanced classes, fair learning for all classes
• Label Smoothing - Soft labels instead of hard, prevents overconfidence
• Dropout - Random neuron deactivation, regularization technique
• EarlyStopping - Stop when not improving, prevents overfitting
• Batch Processing - Process 32 images at once, memory efficiency
• Softmax - Probability distribution, interpretable outputs
• Confusion Matrix - Error visualization, understand model mistakes''')

doc.add_heading('📝 PROJECT SUMMARY', level=2)
doc.add_paragraph('''
Component       | Technology        | Purpose
----------------|-------------------|---------------------------
Base Model      | MobileNetV2       | Feature extraction
Training        | TensorFlow/Keras  | Deep learning framework
Web App         | Streamlit         | User interface
Data            | PlantVillage      | 15 classes of plant diseases
Accuracy        | ~99%              | Validation performance
Output          | Disease+Treatment | Actionable recommendations

Your project demonstrates:
✅ Transfer Learning
✅ Data Augmentation
✅ Class Imbalance Handling
✅ Model Evaluation (Confusion Matrix, Classification Report)
✅ Web Deployment (Streamlit Cloud)
✅ Real-world Application
''')

# Save document
doc.save('AgroVision_Complete_Documentation.docx')
print("✅ Document created: AgroVision_Complete_Documentation.docx")
